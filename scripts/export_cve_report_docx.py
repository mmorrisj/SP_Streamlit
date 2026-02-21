"""
Convert CVE_MITIGATION_REPORT.md into a formatted Word document.
Outputs: docs/CVE_MITIGATION_REPORT.docx
"""

import io
import re
import sys
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colours ────────────────────────────────────────────────────

NAVY = RGBColor(0x1A, 0x36, 0x5D)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
MID_GREY = RGBColor(0x64, 0x74, 0x8B)
LIGHT_GREY = RGBColor(0x94, 0xA3, 0xB8)
GREEN = RGBColor(0x05, 0x96, 0x69)
RED = RGBColor(0xDC, 0x26, 0x26)
BLUE = RGBColor(0x25, 0x63, 0xEB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEADER_BG = "1A365D"
SUBHEADER_BG = "E2E8F0"
EVEN_ROW_BG = "F8FAFC"


# ── Style helpers ──────────────────────────────────────────────

def _set_font(run, name="Calibri", size=10, bold=False, italic=False,
              color: RGBColor = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _add_heading(doc, text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def _set_cell_shading(cell, hex_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_font(cell, size=8, bold=False, color=None, name="Calibri"):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = name
            run.font.size = Pt(size)
            run.bold = bold
            if color:
                run.font.color.rgb = color


def _add_styled_table(doc, headers, rows, col_widths=None):
    """Add a nicely formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        _set_cell_shading(cell, HEADER_BG)
        _set_cell_font(cell, size=8, bold=True, color=WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            if r_idx % 2 == 0:
                _set_cell_shading(cell, EVEN_ROW_BG)
            _set_cell_font(cell, size=8)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacer
    return table


def _add_para(doc, text, size=10, bold=False, italic=False, color=None,
              space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_code_block(doc, text):
    """Add a monospaced code block."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, name="Consolas", size=8, color=DARK_GREY)
    # Light grey background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F1F5F9")
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)
    p.paragraph_format.space_after = Pt(4)
    return p


# ── Main document builder ─────────────────────────────────────

def build_cve_report():
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── TITLE PAGE ─────────────────────────────────────────────

    for _ in range(5):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_para.add_run("CVE Mitigation Report")
    _set_font(r, size=28, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("SoftPower Analytics Docker Stack")
    _set_font(r, size=16, color=MID_GREY)

    doc.add_paragraph()

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, value in [
        ("Date: ", "February 20, 2026"),
        ("  |  Scanner: ", "Docker Scout v1.18.3"),
        ("  |  Platform: ", "linux/amd64"),
    ]:
        r = meta_para.add_run(label)
        _set_font(r, size=10, bold=True, color=MID_GREY)
        r = meta_para.add_run(value)
        _set_font(r, size=10, color=DARK_GREY)

    doc.add_paragraph()

    # Image inventory table on title page
    _add_styled_table(doc,
        ["Image", "Tag", "Base", "Purpose"],
        [
            ["mmorrisj/softpower-analytics", "1.1.0", "python:3.13-slim (Debian Trixie)",
             "Application (FastAPI + Streamlit + React + ML)"],
            ["mmorrisj/pgvector", "0.8.0-pg16", "postgres:16-bookworm (Debian Bookworm)",
             "PostgreSQL 16 + pgvector extension"],
        ],
        col_widths=[4.5, 2.5, 5, 6],
    )

    doc.add_page_break()

    # ── EXECUTIVE SUMMARY ──────────────────────────────────────

    _add_heading(doc, "Executive Summary", level=1)

    _add_para(doc,
        "A comprehensive vulnerability scan and remediation was performed on both "
        "Docker images in the SoftPower Analytics stack. Combined results:",
        size=10)

    # -- Application Image summary --
    _add_heading(doc, "Application Image (softpower-analytics:1.1.0)", level=2)

    _add_para(doc,
        "59 of 95 identified CVEs were eliminated through base image upgrades, "
        "package version bumps, and build-tool removal. All CRITICAL and HIGH "
        "severity vulnerabilities have been resolved.",
        size=10, bold=True)

    _add_styled_table(doc,
        ["Metric", "Before (v1.0.0)", "After (v1.1.0)", "Change"],
        [
            ["Critical", "2", "0", "-2"],
            ["High", "5", "0", "-5"],
            ["Medium", "6", "1", "-5"],
            ["Low", "81", "35", "-46"],
            ["Unknown", "1", "0", "-1"],
            ["Total", "95", "36", "-59 (62% reduction)"],
            ["Packages Scanned", "418", "370", "-48 removed"],
        ],
        col_widths=[4, 3.5, 3.5, 4],
    )

    _add_para(doc,
        "Risk Assessment: The 36 remaining vulnerabilities are all LOW severity (35) "
        "or MEDIUM (1) with no available upstream fix. None are exploitable in this "
        "application's deployment context. See Section 2 for detailed analysis.",
        size=10, italic=True, color=MID_GREY)

    # -- Database Image summary --
    _add_heading(doc, "Database Image (pgvector:0.8.0-pg16)", level=2)

    _add_para(doc,
        "The previous database image (ankane/pgvector:latest) had 337 CVEs due to "
        "a stale, unmaintained base. A custom image was built from the current "
        "postgres:16-bookworm base with pgvector 0.8.0 compiled from source and "
        "build tools removed.",
        size=10)

    _add_styled_table(doc,
        ["Metric", "Before (ankane/pgvector)", "After (mmorrisj/pgvector)", "Change"],
        [
            ["Critical", "9", "1", "-8"],
            ["High", "96", "6", "-90"],
            ["Medium", "87", "12", "-75"],
            ["Low", "145", "40", "-105"],
            ["Total", "337", "59", "-278 (82% reduction)"],
            ["Packages Scanned", "224", "211", "-13 removed"],
        ],
        col_widths=[4, 4.5, 4.5, 4],
    )

    _add_para(doc,
        "Risk Assessment: The 59 remaining CVEs are all inherited from the official "
        "postgres:16-bookworm base image. The 1 CRITICAL and 6 HIGH are in Go stdlib "
        "(bundled by the PostgreSQL Docker maintainers, not by PostgreSQL itself) and "
        "affect every postgres:16 image on Docker Hub. See Section 5 for detailed analysis.",
        size=10, italic=True, color=MID_GREY)

    doc.add_page_break()

    # ── SECTION 1: REMEDIATION ACTIONS ─────────────────────────

    _add_heading(doc, "1. Remediation Actions Taken", level=1)

    # 1.1 Base Image Upgrade
    _add_heading(doc, "1.1 Base Image Upgrade", level=2)
    _add_styled_table(doc,
        ["Change", "Detail"],
        [
            ["Previous base", "python:3.11-slim (Debian Bookworm)"],
            ["New base", "python:3.13-slim (Debian Trixie)"],
            ["Impact", "Eliminated all Python 3.11-specific CVEs and inherited newer Debian package versions"],
        ],
        col_widths=[4, 14],
    )

    # 1.2 Application Package Upgrades
    _add_heading(doc, "1.2 Application Package Upgrades", level=2)
    _add_styled_table(doc,
        ["Package", "Previous Version", "New Version", "CVEs Fixed"],
        [
            ["PyTorch (torch)", "2.5.1", ">=2.6.0 (CPU-only)", "CVE-2025-32434 (CRITICAL: arbitrary code exec via torch.load)"],
            ["PyArrow", "17.0.0", ">=18.0.0", "CVE-2024-52338 (CRITICAL: arbitrary code exec via IPC deserialization)"],
            ["FastAPI", "0.112.x", ">=0.115.0", "CVE-2024-24762 (HIGH: ReDoS in multipart), CVE-2025-24750 (HIGH)"],
            ["Starlette", "0.37.x", ">=0.40.0 (via FastAPI)", "CVE-2024-47874 (HIGH: multipart DoS), CVE-2025-24750, CVE-2024-24762"],
            ["Pillow", "not pinned", ">=12.1.1", "CVE-2025-3043 (HIGH: buffer overflow in TIFF)"],
            ["wheel", "not pinned", ">=0.46.2", "CVE-2025-33685 (HIGH: path traversal in wheel extraction)"],
            ["setuptools", "70.2.0 (base)", ">=78.1.1", "CVE-2025-47273 (HIGH: path traversal in package install)"],
            ["pip", "24.x (base)", ">=26.0", "CVE-2026-1703 (LOW: path traversal)"],
            ["filelock", "not pinned", ">=3.20.3", "CVE-2025-23207, CVE-2025-23016 (MEDIUM: symlink attacks)"],
            ["pandas", "2.1.4", ">=2.2.0", "Python 3.13 compatibility"],
            ["numpy", "not pinned", ">=2.1.0", "Python 3.13 compatibility"],
            ["psycopg2-binary", "2.9.9", ">=2.9.10", "Python 3.13 wheel availability"],
            ["matplotlib", "not pinned", ">=3.9.0", "Python 3.13 compatibility"],
        ],
        col_widths=[3, 3, 3.5, 8],
    )

    # 1.3 Build Tool Removal
    _add_heading(doc, "1.3 Build Tool Removal (Attack Surface Reduction)", level=2)
    _add_para(doc,
        "build-essential (gcc, g++, binutils, make, dpkg-dev) is required during "
        "pip install for compiling C extensions but is not needed at runtime. "
        "These packages were:",
        size=10)

    p = doc.add_paragraph()
    r = p.add_run("1. ")
    _set_font(r, size=10, bold=True)
    r = p.add_run("Purged via apt-get purge -y build-essential && apt-get autoremove -y")
    _set_font(r, size=10)
    r = p.add_run(" -- removes binaries")
    _set_font(r, size=10, color=MID_GREY)

    p = doc.add_paragraph()
    r = p.add_run("2. ")
    _set_font(r, size=10, bold=True)
    r = p.add_run("Deep-cleaned via dpkg --purge --force-all")
    _set_font(r, size=10)
    r = p.add_run(" -- removes residual package metadata that scanners flag")
    _set_font(r, size=10, color=MID_GREY)

    _add_para(doc,
        "This eliminated 39 binutils-related LOW CVEs and reduced the package "
        "count from 418 to 370.",
        size=10, bold=True)

    # 1.4 Python 3.13 Compatibility
    _add_heading(doc, "1.4 Python 3.13 Compatibility Fixes", level=2)
    _add_para(doc,
        "Application code was updated for Python 3.13 deprecations:",
        size=10)
    _add_styled_table(doc,
        ["File", "Change"],
        [
            ["server/main.py", "datetime.utcnow() -> datetime.now(timezone.utc) (5 locations)"],
            ["server/main.py", "Pydantic regex= -> pattern= (5 locations)"],
            ["server/auth.py", "datetime.utcnow() -> datetime.now(timezone.utc) (2 locations)"],
            ["server/report_validator.py", "datetime.utcnow() -> datetime.now(timezone.utc) (2 locations)"],
        ],
        col_widths=[5, 13],
    )

    doc.add_page_break()

    # ── SECTION 2: REMAINING VULNERABILITIES ───────────────────

    _add_heading(doc, "2. Remaining Vulnerabilities (36 total)", level=1)

    _add_heading(doc, "2.1 Summary", level=2)
    _add_para(doc,
        "All 36 remaining CVEs have no upstream fix available from Debian or PyPI "
        "maintainers (with one exception noted below). They fall into two categories:",
        size=10)
    _add_para(doc, "  \u2022  35 LOW -- Minimal severity, no known active exploitation", size=10)
    _add_para(doc, "  \u2022  1 MEDIUM -- tar vulnerability with limited attack surface", size=10)

    # 2.2 MEDIUM
    _add_heading(doc, "2.2 MEDIUM Severity (1)", level=2)
    _add_styled_table(doc,
        ["CVE", "Package", "Description", "Relevance to Deployment"],
        [
            ["CVE-2025-45582", "tar 1.35",
             "Vulnerability in GNU tar archive handling",
             "Not relevant. tar is not used by the application at runtime. "
             "No user-supplied archives are processed. No fix available from Debian."],
        ],
        col_widths=[3, 2.5, 4.5, 8],
    )

    # 2.3 LOW by package
    _add_heading(doc, "2.3 LOW Severity by Package (35)", level=2)

    # glibc
    _add_heading(doc, "glibc 2.41 -- 7 LOWs", level=3)
    _add_styled_table(doc,
        ["CVE", "Year", "Description", "Relevance"],
        [
            ["CVE-2019-9192", "2019", "Stack exhaustion in regex matching",
             "Not exploitable. Requires attacker-controlled regex patterns."],
            ["CVE-2019-1010025", "2019", "ASLR weakness in __libc_memalign",
             "Not exploitable. Theoretical; requires local code execution."],
            ["CVE-2019-1010024", "2019", "ASLR information leak via /proc",
             "Not exploitable. Requires access to /proc/self/maps."],
            ["CVE-2019-1010023", "2019", "ASLR bypass in __realpath",
             "Not exploitable. Theoretical; requires local code execution."],
            ["CVE-2019-1010022", "2019", "Stack guard bypass",
             "Not exploitable. Theoretical; requires existing code execution."],
            ["CVE-2018-20796", "2018", "Stack exhaustion in regex with backrefs",
             "Not exploitable. Requires attacker-controlled regex."],
            ["CVE-2010-4756", "2010", "glob() resource consumption",
             "Not exploitable. Application does not call glob() on user-supplied paths."],
        ],
        col_widths=[3.5, 1.5, 5, 7],
    )
    _add_para(doc,
        "Assessment: These are long-standing theoretical issues in glibc that Debian "
        "has assessed as LOW impact. All require local code execution or attacker-controlled "
        "inputs. No action needed.",
        size=9, italic=True, color=MID_GREY)

    # openldap
    _add_heading(doc, "openldap 2.6.10 -- 5 LOWs", level=3)
    _add_styled_table(doc,
        ["CVE", "Year", "Description", "Relevance"],
        [
            ["CVE-2026-22185", "2026", "OpenLDAP vulnerability",
             "Not relevant. Application does not use LDAP. libldap is a transitive dependency of curl."],
            ["CVE-2020-15719", "2020", "Certificate hostname validation issue",
             "Not relevant. No LDAP connections made."],
            ["CVE-2017-17740", "2017", "slapd (LDAP server) memory leak",
             "Not relevant. slapd is not installed; only client library present."],
            ["CVE-2017-14159", "2017", "slapd daemon issue",
             "Not relevant. slapd not installed."],
            ["CVE-2015-3276", "2015", "Incorrect certificate DN matching",
             "Not relevant. No LDAP connections."],
        ],
        col_widths=[3.5, 1.5, 5, 7],
    )
    _add_para(doc,
        "Assessment: OpenLDAP client libraries are pulled in as a transitive dependency "
        "of curl/libcurl. The application makes no LDAP connections. No action needed.",
        size=9, italic=True, color=MID_GREY)

    # curl
    _add_heading(doc, "curl 8.14.1 -- 4 LOWs", level=3)
    _add_styled_table(doc,
        ["CVE", "Year", "Description", "Relevance"],
        [
            ["CVE-2025-15224", "2025", "curl vulnerability",
             "Low risk. curl is used only for Docker HEALTHCHECK probes to localhost."],
            ["CVE-2025-15079", "2025", "curl vulnerability",
             "Low risk. Same assessment as above."],
            ["CVE-2025-14017", "2025", "curl vulnerability",
             "Low risk. Same assessment."],
            ["CVE-2025-10966", "2025", "curl vulnerability",
             "Low risk. Same assessment."],
        ],
        col_widths=[3.5, 1.5, 5, 7],
    )
    _add_para(doc,
        "Assessment: curl is installed solely for the container health check "
        "(curl -f http://localhost:8000/api/health). No user-controlled URLs. No action needed.",
        size=9, italic=True, color=MID_GREY)

    # systemd
    _add_heading(doc, "systemd 257.9 -- 4 LOWs", level=3)
    _add_styled_table(doc,
        ["CVE", "Year", "Description", "Relevance"],
        [
            ["CVE-2023-31439", "2023", "systemd-resolved DNS issue",
             "Not relevant. systemd is not the init system in the container."],
            ["CVE-2023-31438", "2023", "systemd-resolved issue",
             "Not relevant. Same as above."],
            ["CVE-2023-31437", "2023", "systemd-resolved issue",
             "Not relevant. Same."],
            ["CVE-2013-4392", "2013", "TOCTOU race condition",
             "Not relevant. Requires running systemd as PID 1; container uses supervisord."],
        ],
        col_widths=[3.5, 1.5, 5, 7],
    )
    _add_para(doc,
        "Assessment: Only libsystemd0 (a shared library) is present. The systemd service "
        "manager and systemd-resolved are not installed or running. Not applicable. No action needed.",
        size=9, italic=True, color=MID_GREY)

    # krb5
    _add_heading(doc, "krb5 (Kerberos) 1.21.3 -- 3 LOWs", level=3)
    _add_styled_table(doc,
        ["CVE", "Year", "Description", "Relevance"],
        [
            ["CVE-2024-26461", "2024", "Memory leak in Kerberos client",
             "Not relevant. Application does not use Kerberos. Library is transitive dependency."],
            ["CVE-2024-26458", "2024", "Memory leak in GSS-API",
             "Not relevant. GSS-API not used."],
            ["CVE-2018-5709", "2018", "Integer overflow in kadmin",
             "Not relevant. kadmin not installed; library only."],
        ],
        col_widths=[3.5, 1.5, 5, 7],
    )
    _add_para(doc,
        "Assessment: Kerberos libraries are transitive dependencies. The application uses "
        "JWT-based authentication, not Kerberos. No action needed.",
        size=9, italic=True, color=MID_GREY)

    # coreutils
    _add_heading(doc, "coreutils 9.7 -- 2 LOWs", level=3)
    _add_styled_table(doc,
        ["CVE", "Year", "Description", "Relevance"],
        [
            ["CVE-2025-5278", "2025", "coreutils vulnerability",
             "Not exploitable. Coreutils commands not invoked on user-supplied input."],
            ["CVE-2017-18018", "2017", "chown following symlinks",
             "Not exploitable. No chown on user-supplied paths."],
        ],
        col_widths=[3.5, 1.5, 5, 7],
    )
    _add_para(doc,
        "Assessment: Standard OS utilities. Not invoked with untrusted input. No action needed.",
        size=9, italic=True, color=MID_GREY)

    # Other packages -- 1 LOW each
    _add_heading(doc, "Other Packages -- 1 LOW each (9 packages)", level=3)
    _add_styled_table(doc,
        ["CVE", "Package", "Description", "Relevance"],
        [
            ["CVE-2021-45346", "sqlite3 3.46.1", "Memory leak via UPDATE",
             "Not relevant. SQLite not used; PostgreSQL is the database."],
            ["CVE-2011-4116", "perl 5.40.1", "Temp file race condition",
             "Not relevant. Perl not invoked by the application."],
            ["CVE-2026-26013", "langchain-core 0.3.83", "SSRF in URL handling",
             "Low risk. Fix requires langchain-core >=1.2.11 (breaking change). CVSS 3.7. See Section 3."],
            ["CVE-2025-15282", "python3.13 3.13.5", "Python vulnerability",
             "Low risk. No fix available from Debian yet."],
            ["CVE-2011-3374", "apt 3.0.3", "Repository signature check issue",
             "Not relevant. apt is not used at runtime."],
            ["CVE-2022-0563", "util-linux 2.41", "chfn/chsh info leak",
             "Not relevant. These utilities are not used."],
            ["CVE-2007-5686", "shadow 4.17.4", "/etc/login.defs info leak",
             "Not relevant. No interactive logins in the container."],
            ["CVE-2010-0928", "openssl 3.5.4", "Theoretical plaintext recovery",
             "Not exploitable. CVSS 2.6. Unlikely in modern TLS."],
            ["CVE-2019-12105", "supervisor 4.2.5", "Information disclosure in web UI",
             "Not exploitable. Supervisor web interface not enabled."],
        ],
        col_widths=[3, 3.5, 4, 7],
    )

    doc.add_page_break()

    # ── SECTION 3: ACTIONABLE ITEMS ────────────────────────────

    _add_heading(doc, "3. Actionable Items for Future Releases", level=1)
    _add_styled_table(doc,
        ["Priority", "Item", "Detail", "Timeline"],
        [
            ["Low", "langchain-core SSRF",
             "CVE-2026-26013 is fixable by upgrading to langchain-core >=1.2.11, "
             "but this is a breaking major version change (0.3.x -> 1.2.x). Current "
             "risk is LOW (CVSS 3.7).",
             "Evaluate during next langchain major version upgrade cycle"],
            ["Monitor", "tar MEDIUM",
             "CVE-2025-45582 -- no fix available. Monitor Debian security tracker. "
             "Not exploitable in current deployment.",
             "Watch for Debian fix"],
            ["Monitor", "curl LOWs",
             "Four unfixed curl CVEs. curl is only used for localhost health checks. "
             "Consider replacing with Python-based probe.",
             "Optional future hardening"],
        ],
        col_widths=[2.5, 3.5, 8, 4],
    )

    # ── SECTION 4: DEPLOYMENT RISK ASSESSMENT ──────────────────

    _add_heading(doc, "4. Deployment Risk Assessment", level=1)
    _add_heading(doc, "Architecture Mitigations", level=2)

    mitigations = [
        ("1. Container Isolation",
         "The application runs in a Docker container with no privileged access. "
         "OS-level CVEs (glibc, systemd, shadow, coreutils) require local code "
         "execution that container boundaries prevent."),
        ("2. No Interactive Shell Access",
         "The container runs supervisord as PID 1 managing FastAPI and Streamlit. "
         "There is no SSH, no login shell, and no interactive user access."),
        ("3. Network Segmentation",
         "The container exposes only ports 8000 (FastAPI API + React UI) and "
         "8501 (Streamlit dashboard). Internal services communicate via localhost only."),
        ("4. Minimal Package Footprint",
         "Build tools (build-essential, binutils, gcc) are removed after compilation. "
         "The final image contains 370 packages -- only what is needed for runtime."),
        ("5. Offline ML Model",
         "The HuggingFace sentence-transformer model is baked into the image at build "
         "time. TRANSFORMERS_OFFLINE=1 and HF_HUB_OFFLINE=1 prevent any runtime network "
         "calls to HuggingFace, eliminating a potential supply-chain vector."),
        ("6. No User-Supplied File Processing",
         "The containerized application serves a web API and dashboard. Document "
         "ingestion runs separately on the host. Users cannot upload arbitrary files "
         "that would trigger the remaining CVEs."),
    ]

    for title, desc in mitigations:
        p = doc.add_paragraph()
        r = p.add_run(title + ": ")
        _set_font(r, size=10, bold=True, color=NAVY)
        r = p.add_run(desc)
        _set_font(r, size=10)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    _add_heading(doc, "Conclusion", level=2)
    _add_para(doc,
        "The softpower-analytics:1.1.0 image is suitable for production deployment. "
        "All CRITICAL and HIGH vulnerabilities have been resolved. The 36 remaining "
        "vulnerabilities (1 MEDIUM, 35 LOW) are either:",
        size=10, bold=True)

    conclusions = [
        "In OS packages not used by the application (systemd, openldap, krb5, sqlite3, perl, shadow)",
        "In utilities used only for internal operations (curl for health checks, tar/apt for build-time only)",
        "Theoretical/academic issues with no practical exploit path in this deployment model",
        "Unfixed upstream with no available patch from Debian or PyPI maintainers",
    ]
    for c in conclusions:
        _add_para(doc, f"  \u2022  {c}", size=10)

    _add_para(doc,
        "No remaining CVE is exploitable through the application's exposed attack "
        "surface (HTTP API on port 8000, Streamlit on port 8501).",
        size=10, bold=True, color=GREEN)

    doc.add_page_break()

    # ── SECTION 5: DATABASE IMAGE ──────────────────────────────

    _add_heading(doc, "5. Database Image: pgvector CVE Analysis", level=1)

    _add_heading(doc, "5.1 Remediation Actions", level=2)
    _add_styled_table(doc,
        ["Action", "Detail"],
        [
            ["Replaced ankane/pgvector:latest",
             "Unmaintained image with 337 CVEs (9C, 96H, 87M, 145L) based on stale debian:12-slim"],
            ["Built custom mmorrisj/pgvector:0.8.0-pg16",
             "Fresh postgres:16-bookworm base, pgvector 0.8.0 compiled from source"],
            ["Build tool removal",
             "build-essential, git, ca-certificates, postgresql-server-dev-16 purged after compilation"],
            ["Residual metadata cleanup",
             "dpkg --purge --force-all on packages in rc state"],
            ["Updated docker-compose.yml",
             "db service now uses mmorrisj/pgvector:0.8.0-pg16"],
        ],
        col_widths=[5, 13],
    )

    _add_heading(doc, "5.2 Remaining Vulnerabilities (59 total)", level=2)
    _add_para(doc,
        "All 59 remaining CVEs originate from the official postgres:16-bookworm base "
        "image. None were introduced by the pgvector extension or the custom build process.",
        size=10)

    _add_heading(doc, "5.2.1 CRITICAL + HIGH: Go stdlib 1.24.6 (1C, 6H)", level=3)
    _add_styled_table(doc,
        ["CVE", "Severity", "Fixed In", "Description"],
        [
            ["CVE-2025-68121", "CRITICAL", "Go 1.24.13", "Go stdlib vulnerability"],
            ["CVE-2025-61729", "HIGH", "Go 1.24.11", "Go stdlib vulnerability"],
            ["CVE-2025-61726", "HIGH", "Go 1.24.12", "Go stdlib vulnerability"],
            ["CVE-2025-61725", "HIGH", "Go 1.24.8", "Go stdlib vulnerability"],
            ["CVE-2025-61723", "HIGH", "Go 1.24.8", "Go stdlib vulnerability"],
            ["CVE-2025-58188", "HIGH", "Go 1.24.8", "Go stdlib vulnerability"],
            ["CVE-2025-58187", "HIGH", "Go 1.24.9", "Go stdlib vulnerability"],
        ],
        col_widths=[3.5, 2.5, 3, 9],
    )

    _add_para(doc,
        "Key context: PostgreSQL is written in C, not Go. The Go stdlib is bundled "
        "into the official postgres Docker image by the PostgreSQL Docker maintainers "
        "for container entrypoint tooling and health check utilities. These CVEs:",
        size=10, bold=True)

    go_points = [
        "Affect every postgres:16 image on Docker Hub -- official, pgvector, third-party. No tag or variant avoids them.",
        "Are not present in bare-metal PostgreSQL installations from RPM/DEB packages, as those do not include Go.",
        "Are not network-exploitable through PostgreSQL's port 5432. They require local code execution within Go-compiled entrypoint tooling.",
        "Will be resolved automatically when the PostgreSQL Docker team rebuilds with Go >=1.24.13.",
    ]
    for point in go_points:
        _add_para(doc, f"  \u2022  {point}", size=10)

    _add_para(doc,
        "Assessment: Accepted risk. Monitor for updated postgres:16-bookworm base image. "
        "Rebuild and push mmorrisj/pgvector when the upstream fix is available.",
        size=9, italic=True, color=MID_GREY)

    _add_heading(doc, "5.2.2 MEDIUM Severity (12)", level=3)
    _add_para(doc,
        "The 12 MEDIUM CVEs are in Debian Bookworm OS packages (glibc, openssl, perl, "
        "libtasn1, etc.) with no upstream fix available. These are the same class of "
        "unfixed OS-level vulnerabilities seen in the application image and are not "
        "exploitable through the PostgreSQL service.",
        size=10)

    _add_heading(doc, "5.2.3 LOW Severity (40)", level=3)
    _add_para(doc,
        "The 40 LOW CVEs are in Debian Bookworm OS packages -- the same packages "
        "(glibc, openldap, curl, systemd, krb5, coreutils, etc.) assessed in "
        "Section 2.3 above. The same relevance assessment applies.",
        size=10)

    _add_heading(doc, "5.3 Alternative Approaches Evaluated", level=2)
    _add_styled_table(doc,
        ["Option", "CVEs", "Tradeoff", "Decision"],
        [
            ["postgres:16-bookworm (current)", "1C 6H 12M 40L",
             "Latest available tag. Debian/glibc-based, maximum compatibility.", "Selected"],
            ["postgres:16-alpine", "1C 6H 11M 1L",
             "Drops 39 LOWs. Same Go issue. Uses musl libc -- collation risk.", "Rejected"],
            ["postgres:17-bookworm", "1C 6H 12M 40L",
             "Same CVE count. Major PG version change requires migration.", "Not needed"],
            ["Host PostgreSQL (bare metal)", "0 Docker CVEs",
             "pgvector not available as pre-built package for CentOS 7.", "Not feasible"],
        ],
        col_widths=[4, 3, 7, 3],
    )

    _add_heading(doc, "5.4 Database Image Deployment Risk Assessment", level=2)
    db_risks = [
        ("Network exposure",
         "PostgreSQL listens on port 5432 within the Docker network. The Go stdlib "
         "CVEs are in container entrypoint tooling, not in the PostgreSQL wire protocol handler."),
        ("No public exposure",
         "The database container is not exposed to the internet. It communicates only "
         "with the application container via the softpower_net Docker network."),
        ("Read-heavy workload",
         "The database serves a read-heavy analytics dashboard. Write operations are "
         "performed by the ingestion pipeline running on the host."),
        ("Container isolation",
         "The database container runs as the postgres user (non-root) with no privileged capabilities."),
    ]
    for title, desc in db_risks:
        p = doc.add_paragraph()
        r = p.add_run(f"{title}: ")
        _set_font(r, size=10, bold=True, color=NAVY)
        r = p.add_run(desc)
        _set_font(r, size=10)
        p.paragraph_format.space_after = Pt(4)

    _add_para(doc,
        "Conclusion: The mmorrisj/pgvector:0.8.0-pg16 image is suitable for production "
        "deployment. The 1 CRITICAL and 6 HIGH vulnerabilities are in Go entrypoint tooling "
        "(not PostgreSQL), are not network-exploitable, and affect every postgres:16 Docker "
        "image universally. They will be resolved by an upstream rebuild.",
        size=10, bold=True, color=GREEN)

    doc.add_page_break()

    # ── APPENDIX A: COMMITS ────────────────────────────────────

    _add_heading(doc, "Appendix A: Commits", level=1)
    _add_styled_table(doc,
        ["Commit", "Description"],
        [
            ["9d2edff", "Fix 52+ Docker CVEs: upgrade base image (3.11->3.13), packages, and remove build tools"],
            ["4c763da", "Bump pandas, numpy, psycopg2-binary for Python 3.13 wheel compatibility"],
            ["03d3661", "Upgrade setuptools/pip and purge binutils metadata to fix 41 more CVEs"],
            ["35f3e04", "Add CVE mitigation report and update deployment scripts"],
            ["aa4dfda", "Add custom pgvector Dockerfile and replace ankane/pgvector in compose"],
        ],
        col_widths=[3, 15],
    )

    # ── APPENDIX B: FILES MODIFIED ─────────────────────────────

    _add_heading(doc, "Appendix B: Files Modified", level=1)
    _add_styled_table(doc,
        ["File", "Changes"],
        [
            ["docker/registry.Dockerfile", "Base image 3.11->3.13, torch>=2.6.0, build-essential purge + dpkg cleanup, setuptools/pip upgrade"],
            ["docker/api.Dockerfile", "Base image 3.11->3.13, build-essential purge"],
            ["docker/api-production.Dockerfile", "Base image 3.11->3.13, build-essential purge"],
            ["docker/dashboard.Dockerfile", "Base image 3.11->3.13"],
            ["docker/airgap.Dockerfile", "Base image 3.11->3.13"],
            ["requirements.txt", "Version bumps for fastapi, uvicorn, torch, pyarrow, psycopg2-binary; added pillow, filelock, wheel"],
            ["requirements-airgap.txt", "Version bumps for all runtime packages; added pillow, filelock, wheel"],
            ["server/main.py", "Python 3.13 compatibility (datetime.utcnow, Pydantic regex->pattern)"],
            ["server/auth.py", "Python 3.13 compatibility (datetime.utcnow)"],
            ["server/report_validator.py", "Python 3.13 compatibility (datetime.utcnow)"],
            ["docker/pgvector.Dockerfile", "New custom pgvector build from postgres:16-bookworm with build tool removal"],
            ["docker-compose.yml", "Updated db service from ankane/pgvector to mmorrisj/pgvector:0.8.0-pg16"],
        ],
        col_widths=[5, 13],
    )

    # ── APPENDIX C: SCANNER OUTPUT ─────────────────────────────

    _add_heading(doc, "Appendix C: Scanner Output", level=1)

    _add_heading(doc, "Application Image", level=2)
    _add_code_block(doc,
        "Target:  mmorrisj/softpower-analytics:1.1.0\n"
        "Digest:  29cfc0fc51de\n"
        "Platform: linux/amd64\n"
        "Size:    1.0 GB\n"
        "Packages: 370\n"
        "\n"
        "Vulnerabilities: 0C 0H 1M 35L (36 total)\n"
        "Base image:      python:3.13-slim (0C 0H 1M 21L)")

    _add_heading(doc, "Database Image", level=2)
    _add_code_block(doc,
        "Target:  mmorrisj/pgvector:0.8.0-pg16\n"
        "Digest:  df4b453d2ebc\n"
        "Platform: linux/amd64\n"
        "Size:    535 MB\n"
        "Packages: 211\n"
        "\n"
        "Vulnerabilities: 1C 6H 12M 40L (59 total)\n"
        "Base image:      postgres:16-bookworm (1C 6H 12M 40L)\n"
        "\n"
        "Previous image (ankane/pgvector:latest):\n"
        "Vulnerabilities: 9C 96H 87M 145L (337 total)")

    _add_heading(doc, "Combined Stack Totals", level=2)
    _add_styled_table(doc,
        ["", "Application", "Database", "Combined"],
        [
            ["Critical:", "0", "1", "1"],
            ["High:", "0", "6", "6"],
            ["Medium:", "1", "12", "13"],
            ["Low:", "35", "40", "75"],
            ["Total:", "36", "59", "95"],
        ],
        col_widths=[3.5, 3.5, 3.5, 3.5],
    )

    _add_para(doc,
        "All Critical/High in database image are from Go stdlib in the "
        "postgres:16-bookworm base, affecting every postgres:16 image on Docker Hub. "
        "Not network-exploitable via PostgreSQL port 5432.",
        size=9, italic=True, color=MID_GREY)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Scanner: Docker Scout v1.18.3  |  Scan date: February 20, 2026")
    _set_font(r, size=9, color=LIGHT_GREY)

    return doc


# ── Entry point ────────────────────────────────────────────────

if __name__ == "__main__":
    project_root = Path(__file__).resolve().parent.parent
    output_path = project_root / "docs" / "CVE_MITIGATION_REPORT.docx"

    doc = build_cve_report()
    doc.save(str(output_path))
    print(f"Word document saved to: {output_path}")
    print(f"Size: {output_path.stat().st_size / 1024:.1f} KB")
