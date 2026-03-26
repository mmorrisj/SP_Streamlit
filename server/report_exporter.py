"""
Word document exporter for Publication reports.
Converts report JSON data into a formatted .docx file using the GAI Summary Template.
"""

import io
import copy
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.ticker import MaxNLocator


# ── Paths ──────────────────────────────────────────────────────

TEMPLATE_PATH = (
    Path(__file__).parent.parent
    / 'services' / 'publication' / 'templates' / 'GAI_Summary_Template.docx'
)

ATOM_ICON_PATH = (
    Path(__file__).parent.parent
    / 'services' / 'publication' / 'templates' / 'atom.png'
)

# ── Colour palette ──────────────────────────────────────────────

CATEGORY_COLORS = {
    'Economic': RGBColor(0x25, 0x63, 0xEB),
    'Diplomacy': RGBColor(0x7C, 0x3A, 0xED),
    'Social': RGBColor(0x05, 0x96, 0x69),
    'Military': RGBColor(0xDC, 0x26, 0x26),
}

CATEGORY_HEX = {
    'Economic': '#2563eb',
    'Diplomacy': '#7c3aed',
    'Social': '#059669',
    'Military': '#dc2626',
}

ENTITY_TYPE_COLORS = {
    'PERSON': '#7c3aed',
    'ORGANIZATION': '#2563eb',
    'COMPANY': '#059669',
    'LOCATION': '#ea580c',
}

TREND_COLORS = ['#2563eb', '#7c3aed', '#059669', '#ea580c', '#dc2626']

# Map placeholder keys to category names in report data
CATEGORY_PLACEHOLDER_MAP = {
    '{{economic_event_section}}': 'Economic',
    '{{diplomatic_event_section}}': 'Diplomacy',
    '{{social_event_section}}': 'Social',
    '{{military_event_section}}': 'Military',
}


# ── Chart helpers ───────────────────────────────────────────────

def _fig_to_bytes(fig) -> io.BytesIO:
    """Render a matplotlib figure to a PNG BytesIO buffer."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_horizontal_bar(data: List[Dict], label_key: str, value_key: str,
                         color: str = '#4a6fa5', title: str = '',
                         max_items: int = 10) -> io.BytesIO:
    """Generic horizontal bar chart."""
    items = data[:max_items]
    if not items:
        return None

    labels = [d[label_key] for d in items][::-1]
    values = [d[value_key] for d in items][::-1]

    fig, ax = plt.subplots(figsize=(5.5, max(2, len(labels) * 0.35)))
    bars = ax.barh(labels, values, color=color, height=0.6)
    ax.set_xlabel('Count', fontsize=8)
    if title:
        ax.set_title(title, fontsize=10, fontweight='bold')
    ax.tick_params(axis='both', labelsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    # Add value labels
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + max(values) * 0.01, bar.get_y() + bar.get_height() / 2,
                str(val), va='center', fontsize=7)

    fig.tight_layout()
    return _fig_to_bytes(fig)


def _make_materiality_histogram(data: List[Dict]) -> io.BytesIO:
    """Materiality score distribution bar chart."""
    filtered = [d for d in data if d['count'] > 0]
    if not filtered:
        return None

    labels = [d['bin'] for d in filtered]
    values = [d['count'] for d in filtered]

    fig, ax = plt.subplots(figsize=(5.5, 2.5))
    ax.bar(labels, values, color='#4a6fa5', width=0.7)
    ax.set_xlabel('Materiality Score Range', fontsize=8)
    ax.set_ylabel('Event Count', fontsize=8)
    ax.set_title('Materiality Score Distribution', fontsize=10, fontweight='bold')
    ax.tick_params(axis='both', labelsize=7)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return _fig_to_bytes(fig)


def _make_materiality_trends(trends: Dict[str, Any]) -> io.BytesIO:
    """Materiality trends line chart."""
    overall = trends.get('overall_series', [])
    if len(overall) < 2:
        return None

    fig, ax = plt.subplots(figsize=(6, 3))

    # Parse months
    overall_months = [datetime.strptime(p['month'], '%Y-%m-%d') for p in overall]
    overall_scores = [p['avg_score'] for p in overall]

    ax.plot(overall_months, overall_scores, color='#1a365d', linewidth=2.5,
            marker='o', markersize=4, label='Overall', zorder=5)

    # Per-recipient lines
    recipient_series = trends.get('recipient_series', {})
    for i, (recip, series) in enumerate(recipient_series.items()):
        months = [datetime.strptime(p['month'], '%Y-%m-%d') for p in series]
        scores = [p['avg_score'] for p in series]
        color = TREND_COLORS[i % len(TREND_COLORS)]
        ax.plot(months, scores, color=color, linewidth=1.5, linestyle='--',
                marker='s', markersize=3, label=recip)

    ax.set_ylim(0, 10)
    ax.set_ylabel('Avg Materiality Score', fontsize=8)
    ax.set_title('Materiality Trends Over Time', fontsize=10, fontweight='bold')
    ax.tick_params(axis='both', labelsize=7)
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%b %y'))
    ax.xaxis.set_major_locator(mdates.MonthLocator())
    ax.legend(fontsize=7, loc='upper left', framealpha=0.9)
    ax.grid(True, alpha=0.3)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.autofmt_xdate(rotation=30)
    fig.tight_layout()
    return _fig_to_bytes(fig)


# ── Hyperlink helper ────────────────────────────────────────────

def _add_hyperlink(paragraph, text: str, url: str,
                   font_size: int = 8, color_hex: str = '2563EB'):
    """Add a clickable hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Font size (half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)

    # Color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color_hex)
    rPr.append(c)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Font name
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ── Style helpers ───────────────────────────────────────────────

def _set_font(run, name='Calibri', size=11, bold=False, italic=False,
              color: RGBColor = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 2,
                 color: RGBColor = None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading


# ── Template placeholder helpers ────────────────────────────────

def _replace_text_in_runs(paragraph, placeholder: str, replacement: str):
    """Replace placeholder text across runs in a paragraph, preserving formatting."""
    # First try simple per-run replacement
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            return True

    # Handle placeholder split across multiple runs
    full_text = ''.join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return False

    # Rebuild: find which runs contain parts of the placeholder
    # Simple approach: concatenate all run text, replace, then redistribute
    runs = paragraph.runs
    if not runs:
        return False

    # Build a mapping of character positions to runs
    combined = ''
    run_boundaries = []
    for run in runs:
        start = len(combined)
        combined += run.text
        run_boundaries.append((start, len(combined), run))

    new_text = combined.replace(placeholder, replacement)

    # Clear all runs and put all text in the first run
    # (preserves first run's formatting)
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ''
    return True


def _replace_in_element_tree(element, placeholder: str, replacement: str):
    """Replace placeholder in all text nodes of an XML element tree."""
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for t_elem in element.iter(f'{ns}t'):
        if t_elem.text and placeholder in t_elem.text:
            t_elem.text = t_elem.text.replace(placeholder, replacement)


def _replace_placeholders_in_doc(doc: Document, placeholder: str, replacement: str):
    """Replace a placeholder in all body paragraphs, headers, footers, and tables."""
    # Body paragraphs
    for paragraph in doc.paragraphs:
        _replace_text_in_runs(paragraph, placeholder, replacement)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_runs(paragraph, placeholder, replacement)

    # Headers and footers
    for section in doc.sections:
        for header_footer in [
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
        ]:
            try:
                for paragraph in header_footer.paragraphs:
                    _replace_text_in_runs(paragraph, placeholder, replacement)
                # Also handle text in drawing/shape elements via XML
                _replace_in_element_tree(header_footer._element, placeholder, replacement)
            except Exception:
                pass


def _find_paragraph_index(doc: Document, placeholder: str) -> int:
    """Find the index of a paragraph containing the given placeholder text."""
    for i, p in enumerate(doc.paragraphs):
        if placeholder in p.text:
            return i
    return -1


def _make_body_paragraph(doc: Document, text: str, style_name: str = 'Body Text - OSE',
                         bold: bool = False, italic: bool = False,
                         font_size: float = None, font_name: str = None,
                         color: RGBColor = None) -> OxmlElement:
    """Create a new paragraph element with the given text and style."""
    p = OxmlElement('w:p')

    # Paragraph properties - apply style
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    # Normalize style name to match internal ID (no spaces, no dashes issue)
    # python-docx stores style IDs without spaces; look up the actual style
    try:
        style = doc.styles[style_name]
        pStyle.set(qn('w:val'), style.style_id)
    except KeyError:
        pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    p.append(pPr)

    # Run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if italic:
        i_elem = OxmlElement('w:i')
        rPr.append(i_elem)
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(szCs)
    if font_name:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
              if isinstance(color, tuple) else str(color))
        rPr.append(c)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)

    p.append(r)
    return p


def _insert_event_content(doc: Document, placeholder: str, cat_data: dict,
                          citations_by_event: List[Dict]):
    """Replace an event section placeholder with event content paragraphs.

    Inserts event paragraphs (name, overview, outcomes) at the position of the
    placeholder paragraph, then removes the placeholder.
    """
    # Find the placeholder paragraph element
    body = doc.element.body
    target_p = None
    for p_elem in body.iterchildren(qn('w:p')):
        # Get full text of paragraph
        texts = [t.text or '' for t in p_elem.iter(qn('w:t'))]
        full_text = ''.join(texts)
        if placeholder in full_text:
            target_p = p_elem
            break

    if target_p is None:
        return

    if not cat_data or not cat_data.get('events'):
        # No events - replace placeholder with "No events reported."
        new_p = _make_body_paragraph(
            doc, 'No events reported for this category.',
            style_name='Body Text - OSE', italic=True
        )
        target_p.addprevious(new_p)
        body.remove(target_p)
        return

    # Build citation lookup for this category
    cat_name = cat_data.get('category', '')
    cat_citations = {}
    for group in citations_by_event:
        if group.get('category') == cat_name:
            for evt_cit in group.get('events', []):
                cat_citations[evt_cit['event_name']] = evt_cit.get('citations', [])

    events = cat_data.get('events', [])
    elements_to_insert = []

    for event in events:
        event_name = event.get('event_name', 'Unnamed Event')
        overview = event.get('overview', '')
        outcomes = event.get('outcomes', '')

        # Event name as bold paragraph
        name_p = _make_body_paragraph(
            doc, event_name,
            style_name='Body Text - OSE', bold=True, font_size=11
        )
        elements_to_insert.append(name_p)

        # Overview
        if overview:
            overview_p = _make_body_paragraph(
                doc, overview, style_name='Body Text - OSE', font_size=10
            )
            elements_to_insert.append(overview_p)

        # Outcomes
        if outcomes:
            outcomes_p = _make_body_paragraph(
                doc, outcomes, style_name='Body Text - OSE', font_size=10
            )
            elements_to_insert.append(outcomes_p)

        # Spacer paragraph
        spacer = _make_body_paragraph(doc, '', style_name='Body Text - OSE')
        elements_to_insert.append(spacer)

    # Insert all elements before the placeholder, then remove placeholder
    for elem in elements_to_insert:
        target_p.addprevious(elem)
    body.remove(target_p)


# ── Main export function (template-based) ──────────────────────

def export_report_to_docx(report_data: dict) -> io.BytesIO:
    """
    Convert a report JSON dict into a formatted .docx BytesIO buffer,
    using the GAI Summary Template.
    """
    doc = Document(str(TEMPLATE_PATH))

    country = report_data.get('country', 'Unknown')
    period_start = report_data.get('period_start', '')
    period_end = report_data.get('period_end', '')
    date_str = f"{_format_date(period_start)} to {_format_date(period_end)}"
    title = report_data.get('title', f'{country} Soft-Power Initiatives')
    overall_summary = report_data.get('overall_summary', '')
    categories = report_data.get('categories', [])
    citations_by_event = report_data.get('citations_by_event', [])
    model_used = report_data.get('model_used', 'GPT-4o')
    generated_at = report_data.get('generated_at', '')

    # Build intro paragraph text
    intro_text = (
        f"During the period from {date_str}, Near Eastern and North African "
        f"media covered the {country}\u2019s soft power initiatives throughout "
        f"the region. OSE leveraged Generative Artificial Intelligence (GAI) "
        f"to summarize key economic, diplomatic, social, military initiatives "
        f"and outcomes covered by Near Eastern media during this timeframe."
    )

    # ── 1. Replace simple placeholders ──────────────────────────
    _replace_placeholders_in_doc(doc, '{{summary_title}}', title)
    _replace_placeholders_in_doc(doc, '{{country}}', country)
    _replace_placeholders_in_doc(doc, '{{date}}', date_str)

    # Replace the intro paragraph (P1) with our generated text
    if len(doc.paragraphs) > 1:
        intro_p = doc.paragraphs[1]
        if '{{date}}' not in intro_p.text and 'During the period' in intro_p.text:
            # Already replaced via placeholder substitution above
            pass
        # If the intro still has template boilerplate, replace it
        if 'Near Eastern and North African media' in intro_p.text:
            # Clear and rewrite with actual summary if available
            if overall_summary:
                for run in intro_p.runs:
                    run.text = ''
                if intro_p.runs:
                    intro_p.runs[0].text = overall_summary
                else:
                    intro_p.add_run(overall_summary)

    # ── 2. Replace event section placeholders ───────────────────
    category_map = {cat['category']: cat for cat in categories}

    for placeholder, cat_name in CATEGORY_PLACEHOLDER_MAP.items():
        cat_data = category_map.get(cat_name)
        _insert_event_content(doc, placeholder, cat_data, citations_by_event)

    # ── 3. Update methodology table ─────────────────────────────
    _update_methodology_table(doc, report_data)

    # ── 4. Add sources page ─────────────────────────────────────
    _add_sources_section(doc, report_data)

    # ── 5. Clean up empty placeholder paragraphs ────────────────
    _cleanup_empty_paragraphs(doc)

    # ── Write to buffer ─────────────────────────────────────────
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _update_methodology_table(doc: Document, report_data: dict):
    """Update the methodology/scope table with report-specific details."""
    if not doc.tables:
        return

    table = doc.tables[0]
    country = report_data.get('country', 'Unknown')
    period_start = report_data.get('period_start', '')
    period_end = report_data.get('period_end', '')
    model_used = report_data.get('model_used', 'GPT-4o')
    date_str = f"{_format_date(period_start)} to {_format_date(period_end)}"

    # The methodology table is in row 1, col 0
    if len(table.rows) < 2:
        return

    cell = table.rows[1].cells[0]

    # Replace placeholders in the existing methodology text
    for paragraph in cell.paragraphs:
        _replace_text_in_runs(paragraph, '{{country}}', country)
        # Replace model reference
        _replace_text_in_runs(paragraph, '[ChatGPT 4.o, 28 February]',
                              f'[{model_used}]')
        # Replace date references
        _replace_text_in_runs(
            paragraph,
            'covering the period from August 2024 to January 2025',
            f'covering the period {date_str}'
        )


def _add_sources_section(doc: Document, report_data: dict):
    """Add sources/citations at the end of the document."""
    citations_by_event = report_data.get('citations_by_event', [])
    if not citations_by_event:
        return

    # Find or create the Sources heading - check if it already exists
    # The template has a "Sources" heading on page 3
    sources_found = False
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'Sources':
            sources_found = True
            # Insert citations after this paragraph
            _insert_citations_after(doc, i, citations_by_event)
            break

    if not sources_found:
        # Add a page break and Sources heading
        doc.add_page_break()
        h = doc.add_paragraph()
        r = h.add_run('Sources')
        _set_font(r, size=14, bold=True)

        _write_citations(doc, citations_by_event)


def _insert_citations_after(doc: Document, para_index: int,
                            citations_by_event: List[Dict]):
    """Insert citation paragraphs after the given paragraph index."""
    body = doc.element.body
    p_elements = list(body.iterchildren(qn('w:p')))

    if para_index >= len(p_elements):
        return

    ref_element = p_elements[para_index]
    elements_to_insert = []

    citation_num = 0
    for group in citations_by_event:
        cat_name = group.get('category', '')

        for event_cit in group.get('events', []):
            event_name = event_cit.get('event_name', '')

            for cit in event_cit.get('citations', []):
                citation_num += 1
                num = cit.get('citation_number', citation_num)
                headline = cit.get('headline', 'Untitled')
                source = cit.get('source_name', 'Unknown')
                pub_date = _format_date(cit.get('published_date'))

                text = f'[{num}] {headline} \u2014 {source}, {pub_date}'

                cit_p = _make_body_paragraph(
                    doc, text, style_name='Body Text - OSE', font_size=9
                )
                elements_to_insert.append(cit_p)

    # Insert after the reference element
    insert_after = ref_element
    for elem in elements_to_insert:
        insert_after.addnext(elem)
        insert_after = elem


def _write_citations(doc: Document, citations_by_event: List[Dict]):
    """Write citations directly to the end of the document."""
    citation_num = 0
    for group in citations_by_event:
        for event_cit in group.get('events', []):
            for cit in event_cit.get('citations', []):
                citation_num += 1
                num = cit.get('citation_number', citation_num)
                headline = cit.get('headline', 'Untitled')
                source = cit.get('source_name', 'Unknown')
                pub_date = _format_date(cit.get('published_date'))
                hyperlink = cit.get('repo_hyperlink', '')

                p = doc.add_paragraph()
                r = p.add_run(f'[{num}] ')
                _set_font(r, size=9, bold=True)
                r = p.add_run(headline)
                _set_font(r, size=9)
                r = p.add_run(f' \u2014 {source}, {pub_date}')
                _set_font(r, size=9, color=RGBColor(0x64, 0x74, 0x8B))
                if hyperlink:
                    r = p.add_run('  ')
                    _add_hyperlink(p, '[ATOM]', hyperlink)


def _cleanup_empty_paragraphs(doc: Document):
    """Remove leftover empty placeholder paragraphs."""
    body = doc.element.body
    for p_elem in list(body.iterchildren(qn('w:p'))):
        texts = [t.text or '' for t in p_elem.iter(qn('w:t'))]
        full_text = ''.join(texts).strip()
        if full_text.startswith('{{') and full_text.endswith('}}'):
            body.remove(p_elem)


# ── Utility ─────────────────────────────────────────────────────

def _format_date(date_str: str | None) -> str:
    if not date_str:
        return ''
    try:
        dt = datetime.strptime(date_str[:10], '%Y-%m-%d')
        return dt.strftime('%b %d, %Y')
    except (ValueError, TypeError):
        return str(date_str)


def _format_datetime(dt_str: str | None) -> str:
    if not dt_str:
        return ''
    try:
        dt = datetime.fromisoformat(dt_str)
        return dt.strftime('%B %d, %Y at %I:%M %p')
    except (ValueError, TypeError):
        return str(dt_str)


# ── Reviewer validation export ─────────────────────────────────

def _extract_citation_numbers(text: str) -> List[int]:
    """Extract all citation numbers like [1], [2], [3] from narrative text."""
    return [int(m) for m in re.findall(r'\[(\d+)\]', text)]


def _build_citation_lookup(citations_by_event: List[Dict]) -> Dict[int, Dict]:
    """Build a flat lookup: citation_number -> citation dict."""
    lookup = {}
    for group in citations_by_event:
        for event_cit in group.get('events', []):
            for cit in event_cit.get('citations', []):
                num = cit.get('citation_number')
                if num is not None:
                    lookup[num] = cit
    return lookup


def _add_image_hyperlink(paragraph, image_path: str, url: str,
                         width: Inches = Inches(0.15)):
    """Add a clickable image (atom.png) as a hyperlink in a paragraph.

    Creates an inline image wrapped in a hyperlink element so clicking
    the icon opens the ATOM repo URL.
    """
    # Create the hyperlink relationship
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )

    # Create the image relationship
    from docx.image.image import Image as DocxImage
    image = DocxImage.from_file(image_path)
    img_r_id = part.relate_to(
        part.package.get_or_add_image_part(image_path).partname,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image',
    )

    # Image dimensions
    width_emu = int(width)
    aspect = image.px_height / image.px_width if image.px_width else 1
    height_emu = int(width_emu * aspect)

    # Build the hyperlink > run > drawing > inline > graphic XML
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')

    drawing = OxmlElement('w:drawing')
    inline = OxmlElement('wp:inline')
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')

    extent = OxmlElement('wp:extent')
    extent.set('cx', str(width_emu))
    extent.set('cy', str(height_emu))
    inline.append(extent)

    effect_extent = OxmlElement('wp:effectExtent')
    effect_extent.set('l', '0')
    effect_extent.set('t', '0')
    effect_extent.set('r', '0')
    effect_extent.set('b', '0')
    inline.append(effect_extent)

    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(id(paragraph) % 100000))
    docPr.set('name', 'ATOM Link')
    docPr.set('descr', 'Open in ATOM')
    inline.append(docPr)

    graphic = OxmlElement('a:graphic')
    graphic.set(qn('xmlns:a'), 'http://schemas.openxmlformats.org/drawingml/2006/main')

    graphic_data = OxmlElement('a:graphicData')
    graphic_data.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')

    ns_pic = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    pic = OxmlElement('pic:pic')
    pic.set(qn('xmlns:pic'), ns_pic)

    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', '0')
    cNvPr.set('name', 'atom.png')
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement('pic:cNvPicPr')
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)

    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), img_r_id)
    blipFill.append(blip)
    stretch = OxmlElement('a:stretch')
    fillRect = OxmlElement('a:fillRect')
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)

    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext = OxmlElement('a:ext')
    ext.set('cx', str(width_emu))
    ext.set('cy', str(height_emu))
    xfrm.append(ext)
    spPr.append(xfrm)
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    pic.append(spPr)

    graphic_data.append(pic)
    graphic.append(graphic_data)
    inline.append(graphic)
    drawing.append(inline)
    run.append(drawing)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_source_line_with_atom(doc: Document, paragraph, cit: Dict,
                               atom_icon_path: str):
    """Add a single source citation line with an ATOM icon hyperlink.

    Format: [N] Headline — Source, Date  [ATOM icon]
    """
    num = cit.get('citation_number', '')
    headline = cit.get('headline', 'Untitled')
    source = cit.get('source_name', 'Unknown')
    pub_date = _format_date(cit.get('published_date'))
    hyperlink = cit.get('repo_hyperlink', '')

    r = paragraph.add_run(f'[{num}] ')
    _set_font(r, size=8, bold=True, color=RGBColor(0x25, 0x63, 0xEB))

    r = paragraph.add_run(headline)
    _set_font(r, size=8)

    r = paragraph.add_run(f' \u2014 {source}, {pub_date}  ')
    _set_font(r, size=8, color=RGBColor(0x64, 0x74, 0x8B))

    if hyperlink:
        _add_hyperlink(paragraph, '[ATOM]', hyperlink, font_size=8,
                       color_hex='2563EB')


def _insert_event_content_reviewer(doc: Document, placeholder: str,
                                   cat_data: dict,
                                   citations_by_event: List[Dict],
                                   citation_lookup: Dict[int, Dict]):
    """Replace event section placeholder with event content + inline source blocks.

    After each overview and outcomes paragraph, inserts a compact source listing
    showing only the citations referenced in that paragraph, with ATOM hyperlinks
    so reviewers can click through to validate each claim.
    """
    body = doc.element.body
    target_p = None
    for p_elem in body.iterchildren(qn('w:p')):
        texts = [t.text or '' for t in p_elem.iter(qn('w:t'))]
        full_text = ''.join(texts)
        if placeholder in full_text:
            target_p = p_elem
            break

    if target_p is None:
        return

    if not cat_data or not cat_data.get('events'):
        new_p = _make_body_paragraph(
            doc, 'No events reported for this category.',
            style_name='Body Text - OSE', italic=True
        )
        target_p.addprevious(new_p)
        body.remove(target_p)
        return

    events = cat_data.get('events', [])

    # We need to build paragraphs using python-docx's high-level API so we
    # can use add_run() etc., then move the underlying XML elements into
    # position.  Trick: append to end of doc, collect elements, reposition.
    created_elements = []

    for event in events:
        event_name = event.get('event_name', 'Unnamed Event')
        overview = event.get('overview', '')
        outcomes = event.get('outcomes', '')

        # ── Event name (bold)
        p = doc.add_paragraph()
        r = p.add_run(event_name)
        _set_font(r, size=11, bold=True)
        try:
            style = doc.styles['Body Text - OSE']
            p.style = style
        except KeyError:
            pass
        created_elements.append(p._p)

        # ── Overview + source block
        if overview:
            p = doc.add_paragraph()
            r = p.add_run(overview)
            _set_font(r, size=10)
            try:
                p.style = doc.styles['Body Text - OSE']
            except KeyError:
                pass
            created_elements.append(p._p)

            # Source block for overview citations
            cited_nums = _extract_citation_numbers(overview)
            if cited_nums:
                src_p = doc.add_paragraph()
                src_p.paragraph_format.left_indent = Pt(18)
                src_p.paragraph_format.space_before = Pt(2)
                src_p.paragraph_format.space_after = Pt(6)
                r = src_p.add_run('Sources: ')
                _set_font(r, size=8, bold=True, color=RGBColor(0x47, 0x55, 0x69))
                src_p.add_run('\n')  # newline before source list

                for cn in cited_nums:
                    cit = citation_lookup.get(cn)
                    if cit:
                        _add_source_line_with_atom(doc, src_p, cit,
                                                   str(ATOM_ICON_PATH))
                        src_p.add_run('\n')
                created_elements.append(src_p._p)

        # ── Outcomes + source block
        if outcomes:
            p = doc.add_paragraph()
            r = p.add_run(outcomes)
            _set_font(r, size=10)
            try:
                p.style = doc.styles['Body Text - OSE']
            except KeyError:
                pass
            created_elements.append(p._p)

            cited_nums = _extract_citation_numbers(outcomes)
            if cited_nums:
                src_p = doc.add_paragraph()
                src_p.paragraph_format.left_indent = Pt(18)
                src_p.paragraph_format.space_before = Pt(2)
                src_p.paragraph_format.space_after = Pt(6)
                r = src_p.add_run('Sources: ')
                _set_font(r, size=8, bold=True, color=RGBColor(0x47, 0x55, 0x69))
                src_p.add_run('\n')

                for cn in cited_nums:
                    cit = citation_lookup.get(cn)
                    if cit:
                        _add_source_line_with_atom(doc, src_p, cit,
                                                   str(ATOM_ICON_PATH))
                        src_p.add_run('\n')
                created_elements.append(src_p._p)

        # Spacer
        spacer = doc.add_paragraph()
        created_elements.append(spacer._p)

    # Move all created elements to before the placeholder, then remove it
    for elem in created_elements:
        target_p.addprevious(elem)
    body.remove(target_p)


def export_reviewer_to_docx(report_data: dict) -> io.BytesIO:
    """
    Export a reviewer validation version of the report.

    Same template-based layout as the publication report, but after each
    overview and outcomes paragraph, a compact source listing is inserted
    showing the cited sources with clickable ATOM hyperlinks so a reviewer
    can quickly open each source document to validate the narrative content.
    """
    doc = Document(str(TEMPLATE_PATH))

    country = report_data.get('country', 'Unknown')
    period_start = report_data.get('period_start', '')
    period_end = report_data.get('period_end', '')
    date_str = f"{_format_date(period_start)} to {_format_date(period_end)}"
    title = report_data.get('title', f'{country} Soft-Power Initiatives')
    overall_summary = report_data.get('overall_summary', '')
    categories = report_data.get('categories', [])
    citations_by_event = report_data.get('citations_by_event', [])
    model_used = report_data.get('model_used', 'GPT-4o')

    # Build flat citation lookup
    citation_lookup = _build_citation_lookup(citations_by_event)

    # ── 1. Replace simple placeholders
    _replace_placeholders_in_doc(doc, '{{summary_title}}',
                                 f'REVIEWER COPY \u2014 {title}')
    _replace_placeholders_in_doc(doc, '{{country}}', country)
    _replace_placeholders_in_doc(doc, '{{date}}', date_str)

    # Replace intro with overall summary
    if len(doc.paragraphs) > 1:
        intro_p = doc.paragraphs[1]
        if 'Near Eastern and North African media' in intro_p.text:
            if overall_summary:
                for run in intro_p.runs:
                    run.text = ''
                if intro_p.runs:
                    intro_p.runs[0].text = overall_summary
                else:
                    intro_p.add_run(overall_summary)

    # ── 2. Replace event section placeholders with reviewer content
    category_map = {cat['category']: cat for cat in categories}

    for placeholder, cat_name in CATEGORY_PLACEHOLDER_MAP.items():
        cat_data = category_map.get(cat_name)
        _insert_event_content_reviewer(doc, placeholder, cat_data,
                                       citations_by_event, citation_lookup)

    # ── 3. Update methodology table
    _update_methodology_table(doc, report_data)

    # ── 4. Add full sources list with ATOM links
    _add_sources_section_reviewer(doc, report_data, citation_lookup)

    # ── 5. Clean up
    _cleanup_empty_paragraphs(doc)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _add_sources_section_reviewer(doc: Document, report_data: dict,
                                  citation_lookup: Dict[int, Dict]):
    """Add the full sources list with ATOM hyperlinks for the reviewer copy."""
    citations_by_event = report_data.get('citations_by_event', [])
    if not citations_by_event:
        return

    # Find the Sources heading in the template
    sources_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'Sources':
            sources_idx = i
            break

    if sources_idx is not None:
        # Write citations after the Sources heading using high-level API,
        # then reposition the XML elements.
        body = doc.element.body
        p_elements = list(body.iterchildren(qn('w:p')))
        ref_element = p_elements[sources_idx]

        insert_after = ref_element
        for group in citations_by_event:
            for event_cit in group.get('events', []):
                for cit in event_cit.get('citations', []):
                    p = doc.add_paragraph()
                    _add_source_line_with_atom(doc, p, cit,
                                               str(ATOM_ICON_PATH))
                    # Move to correct position
                    insert_after.addnext(p._p)
                    insert_after = p._p
    else:
        # No Sources heading - append at end
        doc.add_page_break()
        h = doc.add_paragraph()
        r = h.add_run('Sources')
        _set_font(r, size=14, bold=True)

        for group in citations_by_event:
            for event_cit in group.get('events', []):
                for cit in event_cit.get('citations', []):
                    p = doc.add_paragraph()
                    _add_source_line_with_atom(doc, p, cit,
                                               str(ATOM_ICON_PATH))


# ── Validation status colours ────────────────────────────────────

STATUS_COLORS = {
    'green': RGBColor(0x05, 0x96, 0x69),
    'yellow': RGBColor(0xD9, 0x73, 0x06),
    'red': RGBColor(0xDC, 0x26, 0x26),
}

STATUS_BG_HEX = {
    'green': '#059669',
    'yellow': '#d97306',
    'red': '#dc2626',
}

STATUS_LABELS = {
    'green': 'SUPPORTED',
    'yellow': 'PARTIALLY SUPPORTED',
    'red': 'UNSUPPORTED',
}


def _make_validation_summary_chart(sections: Dict[str, Any]) -> io.BytesIO:
    """Pie chart of green/yellow/red section counts."""
    counts = {'green': 0, 'yellow': 0, 'red': 0}
    for sec in sections.values():
        status = sec.get('status', '').lower()
        if status in counts:
            counts[status] += 1

    labels = []
    sizes = []
    colors = []
    for status in ('green', 'yellow', 'red'):
        if counts[status] > 0:
            labels.append(f"{STATUS_LABELS[status]} ({counts[status]})")
            sizes.append(counts[status])
            colors.append(STATUS_BG_HEX[status])

    if not sizes:
        return None

    fig, ax = plt.subplots(figsize=(4, 3))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, colors=colors, autopct='%1.0f%%',
        startangle=90, textprops={'fontsize': 8}
    )
    for at in autotexts:
        at.set_color('white')
        at.set_fontweight('bold')
    ax.set_title('Validation Results by Section', fontsize=10, fontweight='bold')
    fig.tight_layout()
    return _fig_to_bytes(fig)


def _make_claims_bar_chart(sections: Dict[str, Any]) -> io.BytesIO:
    """Stacked bar chart: supported vs uncited claims per section type."""
    type_counts: Dict[str, Dict[str, int]] = {}
    for sec in sections.values():
        stype = sec.get('section_type', 'unknown')
        if stype not in type_counts:
            type_counts[stype] = {'supported': 0, 'uncited': 0}
        type_counts[stype]['supported'] += sec.get('claims_validated', 0)
        type_counts[stype]['uncited'] += sec.get('uncited_claims', 0)

    if not type_counts:
        return None

    labels = list(type_counts.keys())
    supported = [type_counts[t]['supported'] for t in labels]
    uncited = [type_counts[t]['uncited'] for t in labels]

    fig, ax = plt.subplots(figsize=(5.5, 2.5))
    x = range(len(labels))
    ax.bar(x, supported, label='Supported', color='#059669', width=0.6)
    ax.bar(x, uncited, bottom=supported, label='Uncited', color='#dc2626', width=0.6)
    ax.set_xticks(x)
    ax.set_xticklabels([l.replace('_', ' ').title() for l in labels], fontsize=8)
    ax.set_ylabel('Claims', fontsize=8)
    ax.set_title('Claims by Section Type', fontsize=10, fontweight='bold')
    ax.legend(fontsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    fig.tight_layout()
    return _fig_to_bytes(fig)


def export_validation_to_docx(validation_data: dict,
                               report_title: str = '',
                               country: str = '',
                               period_start: str = '',
                               period_end: str = '',
                               report_data: dict = None) -> io.BytesIO:
    """
    Convert validation results JSON into a formatted .docx BytesIO buffer.

    Args:
        validation_data: Validation JSON dict (status, sections, etc.)
        report_title: Title of the parent report (for context).
        country: Country name.
        period_start: Report period start (YYYY-MM-DD).
        period_end: Report period end (YYYY-MM-DD).
        report_data: Full report JSON dict — used to look up event names
            by index so the DOCX shows real names instead of "Event 3".
    """
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    overall_status = validation_data.get('status', 'unknown').lower()
    sections = validation_data.get('sections', {})
    validated_at = validation_data.get('validated_at', '')
    model_used = validation_data.get('model_used', 'unknown')

    # Build event name lookup: "event:Economic:3" -> "Sailun Group Tire Factory"
    event_name_map: Dict[str, str] = {}
    if report_data:
        for cat in report_data.get('categories', []):
            cat_name = cat.get('category', '')
            for idx, evt in enumerate(cat.get('events', [])):
                key = f"event:{cat_name}:{idx}"
                event_name_map[key] = evt.get('event_name', '')

    # ── Title Page ────────────────────────────────────────────────

    for _ in range(4):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('Source Validation Report')
    _set_font(title_run, size=24, bold=True, color=RGBColor(0x1A, 0x36, 0x5D))

    if report_title:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub_para.add_run(report_title)
        _set_font(r, size=14, color=RGBColor(0x64, 0x74, 0x8B))

    if period_start and period_end:
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = date_para.add_run(
            f"{_format_date(period_start)} — {_format_date(period_end)}"
        )
        _set_font(r, size=12, color=RGBColor(0x64, 0x74, 0x8B))

    # Overall status badge
    status_para = doc.add_paragraph()
    status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # spacer
    badge_para = doc.add_paragraph()
    badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_label = STATUS_LABELS.get(overall_status, overall_status.upper())
    r = badge_para.add_run(f"Overall: {badge_label}")
    _set_font(r, size=16, bold=True,
              color=STATUS_COLORS.get(overall_status, RGBColor(0x33, 0x33, 0x33)))

    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta_para.add_run(
        f"Model: {model_used}  |  "
        f"Validated: {_format_datetime(validated_at)}"
    )
    _set_font(r, size=9, color=RGBColor(0x94, 0xA3, 0xB8))

    doc.add_page_break()

    # ── Summary Statistics ────────────────────────────────────────

    _add_heading(doc, 'Validation Summary', level=1,
                 color=RGBColor(0x1A, 0x36, 0x5D))

    # Compute aggregate stats
    total_sections = len(sections)
    green_count = sum(1 for s in sections.values() if s.get('status', '').lower() == 'green')
    yellow_count = sum(1 for s in sections.values() if s.get('status', '').lower() == 'yellow')
    red_count = sum(1 for s in sections.values() if s.get('status', '').lower() == 'red')
    total_claims = sum(s.get('claims_validated', 0) for s in sections.values())
    total_uncited = sum(s.get('uncited_claims', 0) for s in sections.values())
    total_issues = sum(len(s.get('issues', [])) for s in sections.values())

    stats_para = doc.add_paragraph()
    r = stats_para.add_run(
        f"Sections Validated: {total_sections}  |  "
        f"Total Claims: {total_claims + total_uncited}  |  "
        f"Issues Found: {total_issues}"
    )
    _set_font(r, size=10, bold=True, color=RGBColor(0x33, 0x33, 0x33))

    # Status breakdown line
    breakdown_para = doc.add_paragraph()

    r = breakdown_para.add_run(f"Supported: {green_count}  ")
    _set_font(r, size=10, bold=True, color=STATUS_COLORS['green'])
    r = breakdown_para.add_run(f"Partial: {yellow_count}  ")
    _set_font(r, size=10, bold=True, color=STATUS_COLORS['yellow'])
    r = breakdown_para.add_run(f"Unsupported: {red_count}")
    _set_font(r, size=10, bold=True, color=STATUS_COLORS['red'])

    # Charts
    pie_buf = _make_validation_summary_chart(sections)
    if pie_buf:
        doc.add_picture(pie_buf, width=Inches(3.5))

    bar_buf = _make_claims_bar_chart(sections)
    if bar_buf:
        doc.add_picture(bar_buf, width=Inches(5))

    doc.add_page_break()

    # ── Section-by-Section Details ────────────────────────────────

    _add_heading(doc, 'Detailed Results', level=1,
                 color=RGBColor(0x1A, 0x36, 0x5D))

    # Group sections by type for organized output
    section_groups: Dict[str, List] = {}
    for section_id, section in sections.items():
        stype = section.get('section_type', 'unknown')
        if stype not in section_groups:
            section_groups[stype] = []
        section_groups[stype].append((section_id, section))

    # Define display order
    type_order = ['overall_summary', 'category', 'event', 'entity']
    type_labels = {
        'overall_summary': 'Executive Summary',
        'category': 'Category Narratives',
        'event': 'Event Narratives',
        'entity': 'Entity Summaries',
    }

    for stype in type_order:
        group = section_groups.get(stype)
        if not group:
            continue

        group_label = type_labels.get(stype, stype.replace('_', ' ').title())
        _add_heading(doc, group_label, level=2, color=RGBColor(0x1A, 0x36, 0x5D))

        for section_id, section in group:
            status = section.get('status', 'unknown').lower()
            claims = section.get('claims_validated', 0)
            uncited = section.get('uncited_claims', 0)
            issues = section.get('issues', [])
            summary = section.get('summary', '')

            # Section title with status — use event name map for real names
            display_id = _format_section_id(section_id, event_name_map)
            p = doc.add_paragraph()
            status_label = STATUS_LABELS.get(status, status.upper())
            r = p.add_run(f"[{status_label}]  ")
            _set_font(r, size=9, bold=True,
                      color=STATUS_COLORS.get(status, RGBColor(0x33, 0x33, 0x33)))
            r = p.add_run(display_id)
            _set_font(r, size=10, bold=True)

            # Claims stats
            stats_p = doc.add_paragraph()
            r = stats_p.add_run(
                f"Claims validated: {claims}  |  "
                f"Uncited: {uncited}  |  "
                f"Issues: {len(issues)}"
            )
            _set_font(r, size=9, color=RGBColor(0x64, 0x74, 0x8B))

            if summary:
                sum_p = doc.add_paragraph()
                r = sum_p.add_run(summary)
                _set_font(r, size=9, italic=True, color=RGBColor(0x64, 0x74, 0x8B))

            # Explanation for uncited-only sections (no issues means the
            # narrative contained claims but none had inline citations)
            if uncited > 0 and claims == 0 and not issues:
                note_p = doc.add_paragraph()
                note_p.paragraph_format.left_indent = Pt(18)
                r = note_p.add_run(
                    f"This section contains {uncited} factual claim(s) "
                    f"with no inline citations. The narrative makes "
                    f"assertions that are not attributed to any source "
                    f"document, so they could not be verified."
                )
                _set_font(r, size=9, italic=True,
                          color=STATUS_COLORS['red'])

            # Issues list
            if issues:
                for issue in issues:
                    issue_p = doc.add_paragraph()
                    issue_p.paragraph_format.left_indent = Pt(18)
                    r = issue_p.add_run("- ")
                    _set_font(r, size=9, bold=True, color=RGBColor(0x94, 0xA3, 0xB8))
                    r = issue_p.add_run(issue)
                    _set_font(r, size=9)

            # Spacer
            doc.add_paragraph()

    # ── Write to buffer ──────────────────────────────────────────

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _format_section_id(section_id: str,
                       event_name_map: Dict[str, str] = None) -> str:
    """Convert section_id like 'event:Economic:3' or 'entity:person:Xi Jinping'
    to a readable label.

    Args:
        section_id: Raw section ID from validation JSON.
        event_name_map: Optional dict mapping 'event:Category:idx' to
            the actual event name from the report JSON.
    """
    if section_id == 'overall_summary':
        return 'Overall Summary'

    parts = section_id.split(':')
    if len(parts) == 1:
        return section_id.replace('_', ' ').title()

    stype = parts[0]
    if stype == 'category' and len(parts) >= 2:
        return f"{parts[1]} (Category Narrative)"
    elif stype == 'event' and len(parts) >= 3:
        # Look up event name from report data first
        if event_name_map and section_id in event_name_map:
            return f"{parts[1]} — {event_name_map[section_id]}"
        # Fallback: show index
        try:
            return f"{parts[1]} — Event {int(parts[2]) + 1}"
        except ValueError:
            return f"{parts[1]} — {parts[2]}"
    elif stype == 'entity' and len(parts) >= 3:
        # Entity IDs use names: "entity:person:Xi Jinping"
        entity_type = parts[1].title()
        entity_name = ':'.join(parts[2:])  # Handle names with colons
        return f"{entity_name} ({entity_type})"
    else:
        return section_id.replace(':', ' / ').replace('_', ' ').title()
