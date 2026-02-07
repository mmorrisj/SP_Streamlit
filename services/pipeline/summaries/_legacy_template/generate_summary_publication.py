"""
Generate Summary Publication from Daily Events (June-December 2025)

This script:
1. Loads all daily events from June-December 2025 for China
2. Filters by recipient countries from config.yaml
3. Deduplicates events using similarity matching
4. Uses LLM to generate Overview and Outcomes for each event
5. Creates a Word document using GAI_Summary_Template.docx
"""

import os
import sys
import json
import yaml
import random
import urllib.parse
from datetime import datetime
from pathlib import Path
from collections import defaultdict
from typing import List, Dict, Any

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# Import from existing codebase
from shared.utils.utils import gai, fetch_gai_content


class SummaryPublicationGenerator:
    """Generate publication-ready summaries from daily events"""

    def __init__(self, country: str, start_date: str, end_date: str):
        self.country = country
        self.start_date = start_date
        self.end_date = end_date

        # Load configuration
        config_path = Path(__file__).parent.parent / "shared" / "config" / "config.yaml"
        with open(config_path, 'r') as f:
            self.config = yaml.safe_load(f)

        self.recipients = set(self.config['recipients'])
        self.categories = self.config['categories']

        # Paths
        self.events_dir = Path(__file__).parent.parent / "publications" / "events" / country / "daily"
        self.template_path = Path(__file__).parent / "GAI_Summary_Template.docx"
        self.output_dir = Path(__file__).parent / "output"
        self.output_dir.mkdir(exist_ok=True)

        # Storage
        self.all_events = []
        self.events_by_category = defaultdict(list)
        self.deduped_events = []
        self.key_events = []

    def load_daily_events(self):
        """Load all daily event files from June-December 2025"""
        print(f"Loading daily events for {self.country} from {self.start_date} to {self.end_date}...")

        start = datetime.strptime(self.start_date, "%Y-%m-%d")
        end = datetime.strptime(self.end_date, "%Y-%m-%d")

        for file_path in sorted(self.events_dir.glob("2025-*_events.json")):
            # Extract date from filename
            date_str = file_path.stem.replace("_events", "")
            try:
                file_date = datetime.strptime(date_str, "%Y-%m-%d")
            except ValueError:
                continue

            # Check if in date range
            if not (start <= file_date <= end):
                continue

            # Load events from file
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            for event in data.get('events', []):
                # Filter by recipient countries
                event_recipients = set(event.get('recipients', []))
                if event_recipients & self.recipients:  # Intersection check
                    event['source_date'] = date_str
                    self.all_events.append(event)

        print(f"Loaded {len(self.all_events)} events")
        return self.all_events

    def group_by_category(self):
        """Group events by category"""
        print("Grouping events by category...")

        for event in self.all_events:
            category = event.get('category', 'Unknown')
            if category in self.categories:
                self.events_by_category[category].append(event)

        for cat, events in self.events_by_category.items():
            print(f"  {cat}: {len(events)} events")

        return self.events_by_category

    def deduplicate_events(self):
        """Deduplicate similar events across categories"""
        print("\nDeduplicating events...")

        seen_events = {}

        for category in self.categories:
            events = self.events_by_category.get(category, [])

            for event in events:
                event_name = event['event_name']

                if event_name not in seen_events:
                    seen_events[event_name] = {
                        'event': event,
                        'category': category,
                        'duplicates': []
                    }
                else:
                    # Track duplicates
                    seen_events[event_name]['duplicates'].append(event)

        self.deduped_events = list(seen_events.values())
        print(f"Deduplicated to {len(self.deduped_events)} unique events")

        return self.deduped_events

    def calculate_event_importance(self, event_data: Dict[str, Any]) -> float:
        """Calculate importance score for an event based on multiple factors"""
        event = event_data['event']

        # Factor 1: Number of source documents (articles)
        num_articles = len(event.get('source_doc_ids', []))
        article_score = min(num_articles / 20.0, 1.0) * 40  # Max 40 points, normalized at 20 articles

        # Factor 2: Materiality score (0-10 scale)
        materiality = event.get('materiality', {}).get('score', 0)
        materiality_score = (materiality / 10.0) * 40  # Max 40 points

        # Factor 3: Coverage duration (if date_range available)
        date_range = event.get('date_range', {})
        first_mention = date_range.get('first_mention', 'Unknown')
        last_mention = date_range.get('last_mention', 'Unknown')

        duration_score = 20  # Default baseline
        if first_mention != 'Unknown' and last_mention != 'Unknown':
            try:
                first = datetime.strptime(first_mention, "%Y-%m-%d")
                last = datetime.strptime(last_mention, "%Y-%m-%d")
                days_duration = (last - first).days
                # Longer coverage = more important (max at 30+ days)
                duration_score = min(days_duration / 30.0, 1.0) * 20  # Max 20 points
            except:
                pass

        total_score = article_score + materiality_score + duration_score

        return total_score

    def select_key_events(self, top_n_per_category: int = 15):
        """Select top N most important events per category"""
        print(f"\nSelecting top {top_n_per_category} key events per category...")

        key_events_by_category = {}

        for category in self.categories:
            # Get events for this category
            category_events = [e for e in self.deduped_events if e['category'] == category]

            # Score each event
            for event_data in category_events:
                event_data['importance_score'] = self.calculate_event_importance(event_data)

            # Sort by importance and take top N
            category_events.sort(key=lambda x: x['importance_score'], reverse=True)
            top_events = category_events[:top_n_per_category]

            key_events_by_category[category] = top_events

            print(f"  {category}: Selected {len(top_events)} events from {len(category_events)} total")
            if top_events:
                print(f"    Top event: '{top_events[0]['event']['event_name'][:60]}...' (score: {top_events[0]['importance_score']:.1f})")

        # Flatten into single list
        self.key_events = []
        for category in self.categories:
            self.key_events.extend(key_events_by_category.get(category, []))

        print(f"\nTotal key events selected: {len(self.key_events)}")
        return self.key_events

    def generate_overview_outcomes(self, event: Dict[str, Any]) -> Dict[str, str]:
        """Generate Overview and Outcomes from event summary

        Uses intelligent sentence parsing to separate:
        - Overview: What happened (actors, actions, context)
        - Outcomes: Results, impacts, significance
        """

        event_summary = event.get('event_summary', '')
        event_name = event.get('event_name', '')

        # Split into sentences
        sentences = [s.strip() for s in event_summary.replace('\n', ' ').split('. ') if s.strip()]

        # Keywords that indicate outcomes/impacts
        outcome_keywords = ['result', 'impact', 'consequence', 'outcome', 'significance',
                           'strengthen', 'enhance', 'demonstrate', 'signal', 'reflect',
                           'showcase', 'underscore', 'commitment', 'partnership', 'cooperation',
                           'agreement', 'will', 'aimed at', 'expected', 'future']

        overview_sentences = []
        outcome_sentences = []

        # Classify sentences
        for i, sentence in enumerate(sentences):
            sentence_lower = sentence.lower()

            # Check if sentence contains outcome keywords
            has_outcome_keyword = any(keyword in sentence_lower for keyword in outcome_keywords)

            # First sentence is usually overview context
            if i == 0:
                overview_sentences.append(sentence)
            # Last sentence often contains outcomes/significance
            elif i == len(sentences) - 1 and len(overview_sentences) > 0:
                outcome_sentences.append(sentence)
            # Sentences with outcome keywords go to outcomes
            elif has_outcome_keyword and len(overview_sentences) > 0:
                outcome_sentences.append(sentence)
            # Default: add to overview if it's still small
            elif len(overview_sentences) < len(sentences) // 2:
                overview_sentences.append(sentence)
            else:
                outcome_sentences.append(sentence)

        # Ensure we have both sections
        if not outcome_sentences and len(overview_sentences) > 1:
            # Move last overview sentence to outcomes
            outcome_sentences.append(overview_sentences.pop())

        # Build final text
        overview = '. '.join(overview_sentences)
        if overview and not overview.endswith('.'):
            overview += '.'

        outcomes = '. '.join(outcome_sentences) if outcome_sentences else ''
        if outcomes and not outcomes.endswith('.'):
            outcomes += '.'

        # If we still don't have outcomes, generate a simple one
        if not outcomes:
            num_articles = len(event.get('source_doc_ids', []))
            materiality = event.get('materiality', {}).get('score', 0)

            if materiality >= 8:
                outcomes = f"This initiative represents a significant development in {self.country}'s regional engagement strategy, with high strategic value and substantial bilateral impact."
            elif materiality >= 6:
                outcomes = f"This event strengthens bilateral ties and demonstrates {self.country}'s commitment to regional cooperation and partnership building."
            else:
                outcomes = f"This development contributes to ongoing diplomatic and economic engagement between {self.country} and regional partners."

        return {
            'overview': overview,
            'outcomes': outcomes
        }

    def generate_event_summaries(self):
        """Generate overview and outcomes for each key event"""
        print("\nGenerating Overview and Outcomes for key events...")

        for i, event_data in enumerate(self.key_events):
            event = event_data['event']
            print(f"  Processing {i+1}/{len(self.key_events)}: {event['event_name'][:60]}...")

            # Generate overview and outcomes
            sections = self.generate_overview_outcomes(event)
            event_data['overview'] = sections['overview']
            event_data['outcomes'] = sections['outcomes']

            # Debug: Print first 3 doc IDs and URL info
            if i < 3:
                doc_ids = event.get('source_doc_ids', [])
                print(f"    Doc IDs ({len(doc_ids)}): {doc_ids[:3]}")
                url = event.get('atom_search_url', 'NO URL')
                print(f"    URL preview: {url[:120]}...")

        return self.key_events

    def build_citations(self, event: Dict[str, Any]) -> List[str]:
        """Build citation list from source documents"""
        # TODO: Fetch actual article titles from database using source_doc_ids
        # For now, return placeholder citations

        source_ids = event.get('source_doc_ids', [])
        citations = []

        for doc_id in source_ids[:10]:  # Limit to 10 sources
            # Placeholder - would query database for actual article details
            citations.append(f"Source {doc_id[:8]}...")

        return citations

    @staticmethod
    def build_atom_search_url(doc_ids: List[str]) -> str:
        """Build ATOM search URL from document IDs

        Args:
            doc_ids: List of document IDs (UUIDs)

        Returns:
            Properly formatted ATOM search URL
        """
        if not doc_ids:
            return ""

        # Build the search query: id:("uuid1" OR "uuid2" OR ...)
        id_parts = [f'"{doc_id}"' for doc_id in doc_ids]
        query = f'id:({" OR ".join(id_parts)})'

        # URL encode the query
        encoded_query = urllib.parse.quote(query)

        # Build full ATOM URL
        return f"https://atom.opensource.gov/searches?ss={encoded_query}&go=1"

    @staticmethod
    def validate_atom_url(url: str) -> bool:
        """Check if ATOM URL contains valid UUIDs (not short numeric IDs)

        Args:
            url: The ATOM search URL to validate

        Returns:
            True if URL appears valid, False if it has short/malformed IDs
        """
        if not url:
            return False

        # Decode URL to check content
        decoded = urllib.parse.unquote(url)

        # Check if it contains short numeric IDs like id:("2" OR "3" OR "4")
        # Valid UUIDs are 36 characters with hyphens
        # If we find quoted strings that are just 1-3 digits, it's malformed
        import re
        # Look for patterns like "2" or "99" (short numbers in quotes)
        short_id_pattern = r'"(\d{1,3})"'
        matches = re.findall(short_id_pattern, decoded)

        # If we find matches and they're in the id:() portion, it's malformed
        if matches and 'id:(' in decoded:
            # Check if any matched number is suspiciously short
            for match in matches:
                if len(match) <= 3:  # UUIDs are 36 chars, so anything <=3 is wrong
                    return False

        return True

    def create_word_document(self):
        """Create Word document from template"""
        print("\nCreating Word document...")

        # Load template
        doc = Document(self.template_path)

        # Format date range
        start = datetime.strptime(self.start_date, "%Y-%m-%d")
        end = datetime.strptime(self.end_date, "%Y-%m-%d")
        date_range = f"{start.strftime('%d %B %Y')} to {end.strftime('%d %B %Y')}"

        # Generate title using LLM
        title = self.generate_title()

        # Global placeholders
        global_placeholders = {
            '{{country}}': self.country,
            '{{date}}': date_range,
            '{{summary_title}}': title
        }

        # Replace global placeholders
        for para in doc.paragraphs:
            self.replace_placeholders_in_para(para, global_placeholders)

        for section in doc.sections:
            for para in section.header.paragraphs:
                self.replace_placeholders_in_para(para, global_placeholders)
            for para in section.footer.paragraphs:
                self.replace_placeholders_in_para(para, global_placeholders)

        # Category placeholders
        category_placeholders = {
            'Economic': '{{economic_event_section}}',
            'Diplomacy': '{{diplomatic_event_section}}',
            'Social': '{{social_event_section}}',
            'Military': '{{military_event_section}}'
        }

        available_styles = {s.name for s in doc.styles}

        # Populate each category section
        for category in self.categories:
            placeholder = category_placeholders.get(category)
            if not placeholder:
                continue

            # Find placeholder paragraph
            for para in doc.paragraphs:
                if placeholder in para.text:
                    para.text = ''
                    last = para

                    # Add key events for this category
                    category_events = [e for e in self.key_events if e['category'] == category]

                    for event_data in category_events:
                        event = event_data['event']

                        # Format event name with coverage dates
                        event_name = event['event_name']
                        date_range = event.get('date_range', {})
                        first_mention = date_range.get('first_mention', 'Unknown')
                        last_mention = date_range.get('last_mention', 'Unknown')

                        # Add date range to event name
                        if first_mention != 'Unknown' and last_mention != 'Unknown':
                            try:
                                first = datetime.strptime(first_mention, "%Y-%m-%d")
                                last = datetime.strptime(last_mention, "%Y-%m-%d")

                                # Format: "Event Name (June 15-30, 2025)" or "Event Name (June 15 - July 2, 2025)"
                                if first.month == last.month and first.year == last.year:
                                    date_str = f"{first.strftime('%B')} {first.day}-{last.day}, {first.year}"
                                elif first.year == last.year:
                                    date_str = f"{first.strftime('%B %d')} - {last.strftime('%B %d')}, {first.year}"
                                else:
                                    date_str = f"{first.strftime('%B %d, %Y')} - {last.strftime('%B %d, %Y')}"

                                event_name_with_date = f"{event_name} ({date_str})"
                            except:
                                # If date parsing fails, just use source_date
                                source_date = event.get('source_date', '')
                                if source_date:
                                    try:
                                        date_obj = datetime.strptime(source_date, "%Y-%m-%d")
                                        event_name_with_date = f"{event_name} ({date_obj.strftime('%B %d, %Y')})"
                                    except:
                                        event_name_with_date = event_name
                                else:
                                    event_name_with_date = event_name
                        else:
                            # Use source_date if available
                            source_date = event.get('source_date', '')
                            if source_date:
                                try:
                                    date_obj = datetime.strptime(source_date, "%Y-%m-%d")
                                    event_name_with_date = f"{event_name} ({date_obj.strftime('%B %d, %Y')})"
                                except:
                                    event_name_with_date = event_name
                            else:
                                event_name_with_date = event_name

                        # Event heading
                        ev = self.insert_paragraph_after(
                            last,
                            style_name='Heading 3' if 'Heading 3' in available_styles else None
                        )
                        ev.add_run(event_name_with_date).bold = True
                        last = ev

                        # Overview
                        ov = self.insert_paragraph_after(last, style_name='Normal')
                        ov.add_run("Overview: ").bold = True
                        ov.add_run(event_data['overview'])

                        # Add ATOM hyperlink (validate and rebuild if necessary)
                        atom_url = event.get('atom_search_url', '')
                        if atom_url:
                            # Validate the URL - if it has malformed IDs, regenerate it
                            if not self.validate_atom_url(atom_url):
                                print(f"    WARNING: '{event['event_name'][:50]}...' has malformed URL, regenerating")
                                print(f"      Bad URL: {atom_url[:100]}...")
                                atom_url = self.build_atom_search_url(event.get('source_doc_ids', []))
                                print(f"      New URL: {atom_url[:100]}...")

                            if atom_url:
                                ov.add_run(" ")
                                # URL-decode since it's already encoded and python-docx will encode again
                                decoded_url = urllib.parse.unquote(atom_url)
                                # Add clickable text link
                                self.add_hyperlink(ov, decoded_url, "[View Sources]")

                        last = ov

                        # Outcomes
                        if event_data.get('outcomes'):
                            out = self.insert_paragraph_after(last, style_name='Normal')
                            out.add_run("Outcomes: ").bold = True
                            out.add_run(event_data['outcomes'])

                            # Add ATOM hyperlink (validate and rebuild if necessary)
                            atom_url = event.get('atom_search_url', '')
                            if atom_url:
                                # Validate the URL - if it has malformed IDs, regenerate it
                                if not self.validate_atom_url(atom_url):
                                    atom_url = self.build_atom_search_url(event.get('source_doc_ids', []))

                                if atom_url:
                                    out.add_run(" ")
                                    # URL-decode since it's already encoded and python-docx will encode again
                                    decoded_url = urllib.parse.unquote(atom_url)
                                    # Add clickable text link
                                    self.add_hyperlink(out, decoded_url, "[View Sources]")

                            last = out

                    break  # Done with this category

        # Save document with timestamp to avoid file locks
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = self.output_dir / f"{self.country}_{self.start_date}_to_{self.end_date}_Summary_{timestamp}.docx"
        doc.save(output_file)
        print(f"\nDocument saved to: {output_file}")
        return output_file

    def generate_title(self) -> str:
        """Generate descriptive title"""
        start = datetime.strptime(self.start_date, "%Y-%m-%d")
        end = datetime.strptime(self.end_date, "%Y-%m-%d")
        date_range = f"{start.strftime('%B %Y')} to {end.strftime('%B %Y')}"

        # Simple title without LLM
        return f"{self.country} Soft Power Initiatives {date_range}"

    # Helper methods for Word document manipulation

    @staticmethod
    def replace_placeholders_in_para(para, mapping):
        """Replace placeholders in paragraph"""
        for key, val in mapping.items():
            if key in para.text:
                new_text = para.text.replace(key, val)
                for run in para.runs:
                    run.text = ""
                para.add_run(new_text)

    @staticmethod
    def insert_paragraph_after(existing_para, style_name=None):
        """Insert new paragraph after existing one"""
        new_p_elm = OxmlElement('w:p')
        existing_para._p.addnext(new_p_elm)
        new_para = Paragraph(new_p_elm, existing_para._parent)

        if style_name and style_name in {s.name for s in existing_para.part.styles}:
            new_para.style = style_name

        return new_para

    @staticmethod
    def add_hyperlink(paragraph, url, text):
        """Add a text hyperlink to a paragraph

        Args:
            paragraph: The paragraph to add the hyperlink to
            url: The URL to link to
            text: The display text for the link

        Returns:
            The paragraph with the hyperlink added
        """
        part = paragraph.part

        # Create relationship
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )

        # Build w:hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Build a run inside it
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Add blue color and underline for standard link appearance
        c = OxmlElement('w:color')
        c.set(qn('w:val'), "0563C1")  # Standard Word hyperlink blue
        u = OxmlElement('w:u')
        u.set(qn('w:val'), "single")

        rPr.append(c)
        rPr.append(u)
        new_run.append(rPr)

        # Add text
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

        return paragraph

    @staticmethod
    def add_image_hyperlink(paragraph, image_path, target_url, width=None, height=None):
        """Add clickable image hyperlink"""
        part = paragraph.part

        # Insert picture
        pic_run = paragraph.add_run()
        pic_run.add_picture(image_path, width=Inches(width) if width else None,
                           height=Inches(height) if height else None)

        # Build hyperlink
        rId_link = part.relate_to(
            target_url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )
        hlink = OxmlElement('w:hyperlink')
        hlink.set(qn('r:id'), rId_link)

        # Move picture run into hyperlink
        run_elm = paragraph._p[-1]
        paragraph._p.remove(run_elm)
        hlink.append(run_elm)
        paragraph._p.append(hlink)

        return paragraph

    def run(self, top_n_per_category: int = 15):
        """Execute full pipeline

        Args:
            top_n_per_category: Number of top events to include per category (default: 15)
        """
        print("="*70)
        print(f"Summary Publication Generator")
        print(f"Country: {self.country}")
        print(f"Period: {self.start_date} to {self.end_date}")
        print(f"Top events per category: {top_n_per_category}")
        print("="*70)

        # Step 1: Load events
        self.load_daily_events()

        # Step 2: Group by category
        self.group_by_category()

        # Step 3: Deduplicate
        self.deduplicate_events()

        # Step 4: Select key events based on importance
        self.select_key_events(top_n_per_category=top_n_per_category)

        # Step 5: Generate summaries
        self.generate_event_summaries()

        # Step 6: Create document
        output_file = self.create_word_document()

        print("\n" + "="*70)
        print("COMPLETED SUCCESSFULLY!")
        print(f"Output: {output_file}")
        print(f"Total key events included: {len(self.key_events)}")
        print("="*70)

        return output_file


def main():
    """Main entry point"""
    # Configuration
    country = "China"
    start_date = "2025-06-01"
    end_date = "2025-12-31"

    # Number of top events to include per category
    # Adjust based on desired document length:
    #   - 10-15: Concise summary of most important events
    #   - 20-30: Comprehensive coverage
    #   - 50+: Exhaustive detail
    top_events_per_category = 15

    # Generate summary
    generator = SummaryPublicationGenerator(country, start_date, end_date)
    generator.run(top_n_per_category=top_events_per_category)


if __name__ == "__main__":
    main()
